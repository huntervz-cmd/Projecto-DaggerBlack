import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io

# Modelos de Machine Learning
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score

# Exportación
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image

# Configuración de la página
st.set_page_config(page_title="Sistema de Predicción Académica", layout="wide")

# ==========================================
# 1. CONTROL DE SESIÓN (LOGIN)
# ==========================================
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

def login():
    st.title("🔐 Acceso al Sistema - Panel del Docente")
    st.write("Por favor, introduce tus credenciales para acceder al modelo predictivo y la gestión de alumnos.")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        with st.form("formulario_login"):
            usuario = st.text_input("Usuario (Correo o Cédula)", placeholder="ejemplo@docente.com")
            clave = st.text_input("Contraseña", type="password", placeholder="••••••••")
            boton_ingresar = st.form_submit_button("Iniciar Sesión")
            
            if boton_ingresar:
                # Aquí puedes definir el usuario y clave que desees
                if usuario == "docente2026" and clave == "admin123":
                    st.session_state.autenticado = True
                    st.success("¡Acceso concedido!")
                    st.rerun()
                else:
                    st.error("❌ Usuario o contraseña incorrectos. Inténtalo de nuevo.")

# ==========================================
# 2. INICIALIZACIÓN DE DATOS (ESTADO DE LA APP)
# ==========================================
if 'alumnos' not in st.session_state:
    st.session_state.alumnos = pd.DataFrame({
        'ID': [101, 102, 103, 104, 105],
        'Nombre': ['Xavier Mejias', 'Ana Gómez', 'Carlos Pérez', 'María Silva', 'Luis Martínez'],
        'Asistencia_%': [95, 60, 85, 45, 90],
        'Tareas_Promedio': [9.0, 4.5, 7.5, 3.0, 8.5],
        'Rendimiento_Anterior': [8.5, 5.0, 6.8, 4.0, 7.9],
        'En_Riesgo': [0, 1, 0, 1, 0]
    })

if 'cronograma' not in st.session_state:
    st.session_state.cronograma = pd.DataFrame({
        'ID_Actividad': [1, 2],
        'Actividad': ['Examen Parcial I', 'Entrega de Proyecto'],
        'Fecha': ['2026-06-15', '2026-07-01'],
        'Porcentaje': [30, 20]
    })

# ==========================================
# 3. LÓGICA DE VINCULACIÓN Y ENRUTAMIENTO
# ==========================================
if not st.session_state.autenticado:
    # Si no está logueado, solo se ejecuta la vista de login
    login()
else:
    # Si está logueado, se vincula y muestra toda la aplicación principal
    st.sidebar.title("Bienvenido, Docente")
    
    # Botón para cerrar sesión
    if st.sidebar.button("🚪 Cerrar Sesión"):
        st.session_state.autenticado = False
        st.rerun()
        
    st.sidebar.write("---")
    st.sidebar.title("Navegación del Sistema")
    opcion = st.sidebar.radio("Selecciona una pestaña:", [
        "📊 Dashboard y EDA", 
        "🤖 Modelo Predictivo", 
        "📅 Control de Asistencia", 
        "📆 Cronograma de Actividades"
    ])

    # ==========================================
    # PESTAÑA 1: DASHBOARD Y EDA
    # ==========================================
    if opcion == "📊 Dashboard y EDA":
        st.title("📊 Análisis Exploratorio de Datos (EDA) y Dashboard")
        st.write("Visualiza las correlaciones y el estado actual de los estudiantes.")
        
        df = st.session_state.alumnos
        st.dataframe(df)
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Matriz de Correlación")
            fig, ax = plt.subplots()
            corr = df[['Asistencia_%', 'Tareas_Promedio', 'Rendimiento_Anterior', 'En_Riesgo']].corr()
            cax = ax.matshow(corr, cmap='coolwarm', vmin=-1, vmax=1)
            fig.colorbar(cax)
            ticks = np.arange(0, len(corr.columns), 1)
            ax.set_xticks(ticks)
            ax.set_yticks(ticks)
            ax.set_xticklabels(corr.columns, rotation=45, ha='left')
            ax.set_yticklabels(corr.columns)
            st.pyplot(fig)
            
        with col2:
            st.subheader("Asistencia vs Calificación de Tareas")
            fig2, ax2 = plt.subplots()
            scatter = ax2.scatter(df['Asistencia_%'], df['Tareas_Promedio'], c=df['En_Riesgo'], cmap='bwr', s=100)
            ax2.set_xlabel('Asistencia %')
            ax2.set_ylabel('Promedio de Tareas')
            ax2.axvline(75, color='red', linestyle='--', label='Límite Alerta Asistencia')
            ax2.legend(*scatter.legend_elements(), title="En Riesgo")
            st.pyplot(fig2)

    # ==========================================
    # PESTAÑA 2: MODELO PREDICTIVO
    # ==========================================
    elif opcion == "🤖 Modelo Predictivo":
        st.title("🤖 Entrenamiento de Modelos y Predicción")
        
        df = st.session_state.alumnos
        X = df[['Asistencia_%', 'Tareas_Promedio', 'Rendimiento_Anterior']]
        y = df['En_Riesgo']
        
        X_bounce = pd.concat([X]*5, ignore_index=True)
        y_bounce = pd.concat([y]*5, ignore_index=True)
        
        X_train, X_test, y_train, y_test = train_test_split(X_bounce, y_bounce, test_size=0.3, random_state=42)
        
        algoritmo = st.selectbox("Selecciona el algoritmo de Machine Learning:", ["Logistic Regression", "Decision Tree"])
        
        if algoritmo == "Logistic Regression":
            model = LogisticRegression()
        else:
            model = DecisionTreeClassifier()
            
        model.fit(X_train, y_train)
        preds = model.predict(X_test)
        acc = accuracy_score(y_test, preds)
        
        st.success(f"Modelo {algoritmo} entrenado con una precisión simulada de: {acc * 100:.2f}%")
        
        st.subheader("🔮 Predecir riesgo de un alumno nuevo")
        with st.form("predict_form"):
            input_asistencia = st.slider("Porcentaje de Asistencia", 0, 100, 80)
            input_tareas = st.slider("Promedio de Tareas (0-10)", 0.0, 10.0, 7.0)
            input_anterior = st.slider("Rendimiento Anterior (0-10)", 0.0, 10.0, 6.5)
            
            btn_predict = st.form_submit_button("Calcular Riesgo")
            
            if btn_predict:
                pred_individual = model.predict([[input_asistencia, input_tareas, input_anterior]])
                if pred_individual[0] == 1:
                    st.error("⚠️ Alumno en ALTO RIESGO de deserción o reprobación. Requiere intervención oportuna.")
                else:
                    st.success("✅ Alumno en condición ESTABLE.")

    # ==========================================
    # PESTAÑA 3: CONTROL DE ASISTENCIA (CRUD + EXPORTAR)
    # ==========================================
    elif opcion == "📅 Control de Asistencia":
        st.title("📅 Registro y Control de Asistencia (CRUD)")
        
        df_alumnos = st.session_state.alumnos
        
        st.subheader("Agregar o Editar Alumno")
        col_c1, col_c2, col_c3 = st.columns(3)
        with col_c1:
            id_al = st.number_input("ID Alumno", min_value=100, max_value=999, step=1)
        with col_c2:
            nom_al = st.text_input("Nombre Completo")
        with col_c3:
            asist_al = st.slider("Asistencia Actual %", 0, 100, 85)
            
        col_c4, col_c5 = st.columns(2)
        with col_c4:
            tar_al = st.number_input("Promedio Tareas", min_value=0.0, max_value=10.0, value=7.0)
        with col_c5:
            rend_al = st.number_input("Rendimiento Académico Anterior", min_value=0.0, max_value=10.0, value=7.0)

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("💾 Guardar / Actualizar Alumno"):
                riesgo_calc = 1 if asist_al < 75 or tar_al < 5.0 else 0
                if id_al in df_alumnos['ID'].values:
                    df_alumnos.loc[df_alumnos['ID'] == id_al, ['Nombre', 'Asistencia_%', 'Tareas_Promedio', 'Rendimiento_Anterior', 'En_Riesgo']] = [nom_al, asist_al, tar_al, rend_al, riesgo_calc]
                    st.success("Alumno editado correctamente.")
                else:
                    nuevo_al = pd.DataFrame([[id_al, nom_al, asist_al, tar_al, rend_al, riesgo_calc]], columns=df_alumnos.columns)
                    st.session_state.alumnos = pd.concat([df_alumnos, nuevo_al], ignore_index=True)
                    st.success("Alumno registrado con éxito.")
                st.rerun()

        st.subheader("Eliminar Alumno")
        id_eliminar = st.selectbox("Selecciona ID a eliminar", df_alumnos['ID'].values)
        if st.button("❌ Eliminar Registro"):
            st.session_state.alumnos = df_alumnos[df_alumnos['ID'] != id_eliminar]
            st.warning("Registro eliminado.")
            st.rerun()

        st.write("---")
        st.subheader("Datos Actuales")
        st.dataframe(st.session_state.alumnos)

        st.subheader("💾 Exportar Datos de Asistencia (Respaldo)")
        
        buffer_excel = io.BytesIO()
        st.session_state.alumnos.to_excel(buffer_excel, index=False, engine='openpyxl')
        st.download_button("📥 Descargar Excel", data=buffer_excel.getvalue(), file_name="respaldo_asistencia.xlsx")
        
        doc = Document()
        doc.add_heading('Respaldo de Control de Asistencia', 0)
        for i, row in st.session_state.alumnos.iterrows():
            doc.add_paragraph(f"ID: {row['ID']} | Nombre: {row['Nombre']} | Asistencia: {row['Asistencia_%']}% | Tareas: {row['Tareas_Promedio']}")
        buffer_docx = io.BytesIO()
        doc.save(buffer_docx)
        st.download_button("📥 Descargar Word (DOCX)", data=buffer_docx.getvalue(), file_name="respaldo_asistencia.docx")

        buffer_pdf = io.BytesIO()
        p = canvas.Canvas(buffer_pdf, pagesize=letter)
        p.drawString(100, 750, "Respaldo Estadístico de Alumnos")
        y_pos = 720
        for i, row in st.session_state.alumnos.iterrows():
            p.drawString(100, y_pos, f"ID: {row['ID']} - {row['Nombre']} - Asistencia: {row['Asistencia_%']}%")
            y_pos -= 20
        p.showPage()
        p.save()
        st.download_button("📥 Descargar PDF", data=buffer_pdf.getvalue(), file_name="respaldo_asistencia.pdf")

        fig, ax = plt.subplots(figsize=(6, 3))
        ax.axis('tight')
        ax.axis('off')
        ax.table(cellText=st.session_state.alumnos.values, colLabels=st.session_state.alumnos.columns, loc='center')
        buffer_jpg = io.BytesIO()
        plt.savefig(buffer_jpg, format='jpg', bbox_inches='tight')
        plt.close()
        st.download_button("📥 Descargar Imagen (JPG)", data=buffer_jpg.getvalue(), file_name="tabla_asistencia.jpg", mime="image/jpeg")

    # ==========================================
    # PESTAÑA 4: CRONOGRAMA DE ACTIVIDADES (CRUD)
    # ==========================================
    elif opcion == "📆 Cronograma de Actividades":
        st.title("📆 Cronograma de Actividades")
        df_cron = st.session_state.cronograma
        
        st.dataframe(df_cron)
        
        st.subheader("Agregar / Modificar Actividad")
        col_cr1, col_cr2, col_cr3 = st.columns(3)
        with col_cr1:
            id_act = st.number_input("ID Actividad", min_value=1, max_value=100, step=1)
        with col_cr2:
            nom_act = st.text_input("Nombre de la Actividad")
        with col_cr3:
            fecha_act = st.date_input("Fecha de Entrega")
            
        porc_act = st.slider("Ponderación %", 0, 100, 15)
        
        if st.button("📅 Guardar Actividad"):
            if id_act in df_cron['ID_Actividad'].values:
                df_cron.loc[df_cron['ID_Actividad'] == id_act, ['Actividad', 'Fecha', 'Porcentaje']] = [nom_act, str(fecha_act), porc_act]
            else:
                nueva_act = pd.DataFrame([[id_act, nom_act, str(fecha_act), porc_act]], columns=df_cron.columns)
                st.session_state.cronograma = pd.concat([df_cron, nueva_act], ignore_index=True)
            st.success("Cronograma actualizado.")
            st.rerun()

        st.subheader("Eliminar Actividad")
        id_act_eliminar = st.selectbox("Selecciona ID de actividad a eliminar", df_cron['ID_Actividad'].values)
        if st.button("🗑️ Eliminar Actividad"):
            st.session_state.cronograma = df_cron[df_cron['ID_Actividad'] != id_act_eliminar]
            st.warning("Actividad removida.")
            st.rerun()